from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("output/docx/Jasmine-软件测试报告.docx")
REPORT_DATE = "2026-06-24"
PROJECT_VERSION = "2.0.0"


def set_run_font(run, ascii_font: str = "Calibri", east_asia_font: str = "宋体", size: int | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_layout_fixed(table):
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=10, color="666666")


def apply_base_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.clear()
    label = paragraph.add_run("Jasmine 软件测试报告  第 ")
    set_run_font(label, size=10, color="666666")
    add_page_number(paragraph)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=10, color="666666")


def add_paragraph(doc: Document, text: str, style: str = "Normal", bold_prefix: str | None = None, center: bool = False):
    paragraph = doc.add_paragraph(style=style)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, size=11, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11)
    else:
        run = paragraph.add_run(text)
        size = 11
        if style == "Heading 1":
            size = 16
        elif style == "Heading 2":
            size = 13
        elif style == "Heading 3":
            size = 12
        set_run_font(run, size=size, bold=style.startswith("Heading"))
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15


def populate_table(table, rows, widths, header_fill="D9E2F3", body_font_size=10):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_layout_fixed(table)

    for row_idx, row_values in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx, value in enumerate(row_values):
            cell = row.cells[col_idx]
            cell.text = ""
            set_cell_width(cell, widths[col_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(str(value))
            set_run_font(run, size=body_font_size, bold=(row_idx == 0))
            if row_idx == 0:
                set_cell_shading(cell, header_fill)


def add_summary_box(doc: Document, lines: list[str]):
    table = doc.add_table(rows=len(lines), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_layout_fixed(table)
    for idx, line in enumerate(lines):
        cell = table.cell(idx, 0)
        set_cell_width(cell, 6.5)
        set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "F4F6F9" if idx == 0 else "FAFBFC")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(line)
        set_run_font(run, size=11, bold=(idx == 0), color="1F3A5F" if idx == 0 else None)
    doc.add_paragraph()


def build_document():
    doc = Document()
    apply_base_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(72)
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("Jasmine 软件测试报告")
    set_run_font(run, east_asia_font="黑体", size=24, bold=True, color="1F3A5F")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("测试策略 · 黑盒测试 · 白盒测试")
    set_run_font(run, east_asia_font="黑体", size=14, color="44546A")

    info = doc.add_table(rows=6, cols=2)
    info.style = "Table Grid"
    populate_table(
        info,
        [
            ["项目名称", "Jasmine 桌面音乐播放器"],
            ["文档类型", "基于项目文件的阶段性软件测试报告"],
            ["项目版本", PROJECT_VERSION],
            ["生成日期", REPORT_DATE],
            ["编写依据", "README、package.json、src/ 核心源码、tests/ 自动化测试、graduation-project-notes 缺陷记录"],
            ["报告说明", "本报告以代码与现有自动化验证结果为主，黑盒部分包含测试用例设计与静态核查结论，白盒部分包含模块路径分析与已执行单元测试结果。"],
        ],
        widths=[1.35, 5.15],
        header_fill="F2F4F7",
        body_font_size=10,
    )

    doc.add_paragraph()
    add_summary_box(
        doc,
        [
            "阶段性结论",
            "当前仓库已具备较完整的本地音乐、在线音乐、歌词、歌单、国际化和桌面发布能力，自动化验证链路可正常运行。",
            "2026-06-24 实际执行结果：12 项自动化测试全部通过，TypeScript 构建通过，Vite 生产构建通过。",
            "当前最主要的功能性风险是大歌单在线播放稳定性问题，项目记录中明确指出 2000+ 首歌单存在进度停留 0 和无声音的缺陷，建议作为后续修复优先项。",
        ],
    )

    doc.add_page_break()

    add_paragraph(doc, "1 引言", style="Heading 1")
    add_paragraph(doc, "1.1 文档目的", style="Heading 2")
    add_paragraph(doc, "本文档用于结合 Jasmine 项目的现有代码、配置、自动化测试和缺陷记录，给出一份可直接用于课程作业或毕业设计附件的测试报告。报告重点说明测试范围、测试策略、黑盒测试设计、白盒测试分析、已执行验证结果，以及当前仍需关注的质量风险。")
    add_paragraph(doc, "1.2 项目概况", style="Heading 2")
    add_paragraph(doc, "Jasmine 是一个基于 Vite、React、TypeScript 与 Electron 构建的桌面音乐播放器，支持本地音频导入、网易云在线搜索与播放、扫码登录、歌词展示与编辑、歌单管理、播放历史、分享卡片和中英文界面切换。项目同时包含桌面端打包配置和播放行为统计服务，具有较完整的产品闭环。")
    add_paragraph(doc, "1.3 测试依据", style="Heading 2")
    add_bullet(doc, "功能依据：README.md 中列出的本地音乐、在线音乐、歌词、歌单、历史、国际化、分享与桌面构建能力。")
    add_bullet(doc, "实现依据：src/context/PlayerContext.tsx、src/lib/db.ts、src/lib/neteaseLogin.ts、src/lib/share.ts、analytics/server.cjs 等核心模块。")
    add_bullet(doc, "自动化依据：tests/playback-restore.test.mjs、tests/analytics-api.test.cjs、src/components/discoverArtistHero.test.ts。")
    add_bullet(doc, "风险依据：graduation-project-notes/PLAYBACK_BUG.md 中对大歌单播放异常的排查记录。")

    add_paragraph(doc, "2 测试策略", style="Heading 1")
    add_paragraph(doc, "2.1 测试目标", style="Heading 2")
    add_paragraph(doc, "测试的核心目标是验证 Jasmine 在主要用户链路上的正确性、稳定性与可交付性，即：本地音乐是否可导入并持久化、在线内容是否可搜索与播放、播放器状态是否可恢复、分析服务是否可正确聚合数据、桌面/Web 构建是否可顺利产出。")
    add_paragraph(doc, "2.2 测试方法", style="Heading 2")
    add_bullet(doc, "黑盒测试：从用户视角覆盖导入、搜索、登录、播放、歌单、歌词、历史、分享、国际化等核心功能，采用等价类、边界值和异常场景设计测试用例。")
    add_bullet(doc, "白盒测试：从函数、条件分支和数据流角度分析关键模块，重点检查播放状态恢复、艺人封面提取、统计聚合与播放队列状态流转。")
    add_bullet(doc, "构建验证：通过 TypeScript 编译和 Vite 生产构建验证交付链路，确保项目可以生成发布产物。")
    add_paragraph(doc, "2.3 风险驱动策略", style="Heading 2")
    add_paragraph(doc, "结合项目文件可判断，当前质量风险主要集中在三类场景：第一，在线播放链路依赖本地网易云 API 和网络状态；第二，大型在线歌单属于高边界值场景，已发现稳定性缺陷；第三，PlayerContext 与本地存储的联动较复杂，但现有自动化测试主要覆盖工具函数和分析模块，尚未形成完整的 UI 级回归测试。")

    add_paragraph(doc, "3 测试范围与环境", style="Heading 1")
    add_paragraph(doc, "3.1 测试范围", style="Heading 2")
    scope = doc.add_table(rows=6, cols=4)
    populate_table(
        scope,
        [
            ["模块", "主要功能", "测试关注点", "方法"],
            ["播放器核心", "播放/暂停、上一首/下一首、队列、随机、循环、状态恢复", "状态流转是否正确，边界条件是否可控", "黑盒 + 白盒"],
            ["本地音乐管理", "导入音频、解析元数据、IndexedDB/localStorage 持久化", "导入成功率、元数据完整性、重启恢复", "黑盒 + 代码核查"],
            ["在线音乐模块", "搜索、推荐、扫码登录、歌曲地址、歌词", "接口联通性、登录流程、大歌单边界", "黑盒 + 风险分析"],
            ["辅助体验模块", "歌词显示/编辑、分享卡片、国际化", "界面行为正确、导出可用、语言切换生效", "黑盒"],
            ["分析服务", "播放记录写入、聚合统计、最近记录排序", "插入、汇总、空库边界", "白盒 + 自动化测试"],
        ],
        widths=[1.1, 1.75, 2.55, 1.1],
        body_font_size=9,
    )
    doc.add_paragraph()
    add_paragraph(doc, "3.2 本次验证环境", style="Heading 2")
    env_table = doc.add_table(rows=6, cols=2)
    populate_table(
        env_table,
        [
            ["项目运行栈", "Vite 6、React 18、TypeScript、Electron 41"],
            ["验证日期", REPORT_DATE],
            ["自动化验证命令", "node --test tests/playback-restore.test.mjs tests/analytics-api.test.cjs src/components/discoverArtistHero.test.ts"],
            ["编译验证命令", "node node_modules/typescript/bin/tsc -b"],
            ["构建验证命令", "node node_modules/vite/bin/vite.js build"],
            ["备注", "在线 UI 级联调未在本报告中大规模执行，涉及 API、扫码和音频播放的场景以代码核查及缺陷记录为辅。"],
        ],
        widths=[1.6, 4.9],
        header_fill="F2F4F7",
        body_font_size=10,
    )

    add_paragraph(doc, "4 黑盒测试", style="Heading 1")
    add_paragraph(doc, "4.1 黑盒测试设计说明", style="Heading 2")
    add_paragraph(doc, "黑盒测试围绕用户可感知功能展开，重点选择能够代表完整使用闭环的场景。由于本报告基于项目文件生成，因此表中同时区分“自动化验证通过”“源码核查通过”“已发现缺陷”三类结论，避免将尚未实机联调的场景误写为已执行通过。")
    bb = doc.add_table(rows=11, cols=6)
    populate_table(
        bb,
        [
            ["编号", "测试对象", "输入/操作", "预期结果", "设计方法", "当前结论"],
            ["BB-01", "本地音乐导入", "选择多个合法音频文件导入", "歌曲进入音乐库，元数据被解析并持久化", "等价类", "源码核查通过"],
            ["BB-02", "在线音乐搜索", "输入歌曲名/歌手名发起搜索", "返回歌曲、歌单、艺人等结果", "等价类", "源码核查通过"],
            ["BB-03", "网易云扫码登录", "获取二维码并轮询登录状态", "扫码后返回 cookie 并更新登录状态", "流程测试", "源码核查通过"],
            ["BB-04", "歌单管理", "创建、重命名、删除歌单并添加歌曲", "歌单列表与内容同步更新并保存", "流程测试", "源码核查通过"],
            ["BB-05", "歌词显示与编辑", "打开正在播放页面并查看/编辑 LRC 歌词", "歌词可滚动显示，无歌词时给出提示，编辑后可保存", "等价类 + 异常场景", "源码核查通过"],
            ["BB-06", "播放状态恢复", "应用重开后恢复到上次暂停且有进度的歌曲", "仅在满足恢复条件时自动恢复，不错误恢复到 0 秒新歌", "边界值", "核心判定逻辑自动化通过"],
            ["BB-07", "分享卡片导出", "对当前歌曲生成 PNG 分享卡片", "生成 1080×1080 卡片，封面/标题/歌手信息完整", "等价类", "源码核查通过"],
            ["BB-08", "国际化切换", "在设置页切换中英文界面", "导航、播放控制、设置等文案同步切换", "状态切换", "源码核查通过"],
            ["BB-09", "播放记录分析", "写入多条播放记录并查看汇总", "最近记录按时间倒序，汇总统计正确", "边界值 + 组合测试", "自动化通过"],
            ["BB-10", "大歌单在线播放", "在 2000+ 首网易云歌单中点击播放歌曲", "歌曲应正常获取地址并开始播放", "高边界值", "未通过，已有缺陷记录"],
        ],
        widths=[0.7, 1.05, 1.45, 1.7, 0.8, 0.8],
        body_font_size=7,
    )
    doc.add_paragraph()
    add_paragraph(doc, "4.2 黑盒测试结果分析", style="Heading 2")
    add_paragraph(doc, "从项目文件来看，Jasmine 的基础功能覆盖面较广，且本地导入、扫码登录、分享卡片、国际化和歌词等能力都有明确实现模块支撑，因此大部分功能具备良好的需求可追踪性。当前唯一被项目显式记录为未通过的黑盒场景，是大歌单在线播放边界问题，这也是后续功能验收中必须优先复测的场景。")

    add_paragraph(doc, "5 白盒测试", style="Heading 1")
    add_paragraph(doc, "5.1 白盒测试关注点", style="Heading 2")
    add_paragraph(doc, "白盒测试选择了三类对系统质量影响较大的代码路径：其一是播放状态恢复相关条件分支，其二是封面图片解析与兜底逻辑，其三是分析服务的数据校验、插入、排序和聚合。除此之外，还对 PlayerContext 的状态机和 db.ts 的持久化路径进行了结构性检查，用于识别自动化覆盖不足的部分。")
    wb = doc.add_table(rows=6, cols=5)
    populate_table(
        wb,
        [
            ["编号", "模块/函数", "主要逻辑点", "已覆盖证据", "结论"],
            ["WB-01", "src/lib/playbackRestore.js", "loaded、isPlaying、audioUrl、currentTime 四条件联合判定", "3 个自动化用例全部通过", "分支判断清晰，边界已覆盖"],
            ["WB-02", "src/components/discoverArtistHero.ts", "正则提取 url、渐变背景排除、空值处理、补图条件", "6 个自动化用例全部通过", "典型路径覆盖较完整"],
            ["WB-03", "analytics/server.cjs", "记录写入、最近记录倒序、汇总统计、空库边界", "3 个自动化用例全部通过", "核心数据流稳定"],
            ["WB-04", "src/context/PlayerContext.tsx", "单曲播放回退到单元素队列、NEXT/PREV/REPEAT/SHUFFLE 状态流转", "已做代码路径分析，暂无专门单测", "建议补充 reducer 级测试"],
            ["WB-05", "src/lib/db.ts", "IndexedDB 音频文件存储与 localStorage 数据恢复", "已做代码路径分析，暂无自动化回归", "建议补充浏览器环境测试"],
        ],
        widths=[0.7, 1.35, 2.0, 1.25, 1.2],
        body_font_size=8,
    )
    doc.add_paragraph()
    add_paragraph(doc, "5.2 白盒测试结论", style="Heading 2")
    add_paragraph(doc, "目前仓库内真正形成自动化证据的白盒测试主要集中于工具函数和分析服务，这对于验证边界条件和聚合算法非常有效。但播放器主状态机、数据库持久化和在线播放链路仍然缺少更深层的自动化约束，因此当前白盒测试的优势在于“局部质量较可靠”，不足在于“跨模块协作路径仍需补测”。")

    add_paragraph(doc, "6 已执行验证结果", style="Heading 1")
    results = doc.add_table(rows=4, cols=4)
    populate_table(
        results,
        [
            ["验证项", "执行命令/方式", "结果", "说明"],
            ["自动化测试", "node --test …", "通过", "共 12 项测试，0 失败，覆盖播放恢复、分析服务与艺人封面逻辑"],
            ["TypeScript 编译", "node node_modules/typescript/bin/tsc -b", "通过", "未输出错误信息，说明类型编译成功"],
            ["Vite 生产构建", "node node_modules/vite/bin/vite.js build", "通过", "构建 194 个模块成功，产出 dist/ 资源"],
        ],
        widths=[1.1, 2.3, 0.75, 2.35],
        body_font_size=9,
    )
    doc.add_paragraph()
    add_paragraph(doc, "6.1 结果解释", style="Heading 2")
    add_paragraph(doc, "自动化测试全部通过，说明项目当前已有的关键工具函数与分析服务模块具有较好的可回归性；编译与生产构建通过，说明项目在交付层面不存在明显的阻断性错误。换句话说，仓库当前具备“能编译、能构建、局部核心逻辑可验证”的质量基础。")

    doc.add_page_break()
    add_paragraph(doc, "7 缺陷、风险与建议", style="Heading 1")
    risk_table = doc.add_table(rows=4, cols=4)
    populate_table(
        risk_table,
        [
            ["级别", "问题", "依据", "建议"],
            ["高", "大歌单在线播放在 2000+ 首场景下可能卡在 0 秒且无声音", "graduation-project-notes/PLAYBACK_BUG.md", "优先补充日志、播放链路断点与集成测试，作为验收必测项"],
            ["中", "PlayerContext 状态机路径复杂，但缺少 reducer 级自动化测试", "src/context/PlayerContext.tsx", "围绕 NEXT/PREV/SHUFFLE/REPEAT/QUEUE 补充单元测试"],
            ["中", "本地数据库与浏览器存储路径缺少自动化回归", "src/lib/db.ts", "增加基于 jsdom 或浏览器环境的存储测试"],
        ],
        widths=[0.6, 2.2, 1.4, 2.3],
        body_font_size=9,
    )
    doc.add_paragraph()
    add_paragraph(doc, "7.1 后续测试建议", style="Heading 2")
    add_bullet(doc, "补充 UI 级集成测试，优先覆盖在线搜索->获取音频地址->播放->歌词加载这一主链路。")
    add_bullet(doc, "将大歌单、空歌词、无封面、本地音频损坏文件、登录二维码过期等场景纳入专项回归。")
    add_bullet(doc, "为 PlayerContext 和 db.ts 建立更细粒度的自动化测试，减少跨模块修改带来的回归风险。")
    add_bullet(doc, "在后续版本中考虑引入覆盖率统计，使白盒测试结果更具量化依据。")

    add_paragraph(doc, "8 结论", style="Heading 1")
    add_paragraph(doc, "综合本次基于项目文件的测试分析可以认为：Jasmine 已经具备较完整的软件功能框架，项目代码可通过自动化测试、类型编译和生产构建验证，说明其基本质量处于可交付、可继续迭代的状态。")
    add_paragraph(doc, "与此同时，项目仍存在一个明确的高优先级风险点，即大歌单在线播放稳定性不足；另外，播放器状态机与本地存储链路的自动化覆盖也需要补强。因此，本报告给出的总体评价是：基础能力较完整，局部逻辑验证充分，集成链路仍需继续深化测试。")

    add_paragraph(doc, "附录 A 已执行命令摘要", style="Heading 2")
    add_bullet(doc, "自动化测试：node --test tests/playback-restore.test.mjs tests/analytics-api.test.cjs src/components/discoverArtistHero.test.ts")
    add_bullet(doc, "TypeScript 编译：node node_modules/typescript/bin/tsc -b")
    add_bullet(doc, "Vite 构建：node node_modules/vite/bin/vite.js build")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
