from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


TEMPLATE_PATH = Path("tmp/software-design-template.docx")
OUTPUT_DOCX = Path("output/docx/Jasmine-软件设计规格说明书.docx")
ASSET_DIR = Path("output/docx/assets_sds")
FONT_PATH = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
DATE_TEXT = "2026-06-24"


def font(size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_PATH, size=size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        candidate = current + ch
        if draw.textlength(candidate, font=fnt) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines or [""]


def centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fnt, fill: str):
    lines = text.split("\n")
    bbox = draw.multiline_textbbox((0, 0), "\n".join(lines), font=fnt, spacing=6, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    x = box[0] + (box[2] - box[0] - width) / 2
    y = box[1] + (box[3] - box[1] - height) / 2
    draw.multiline_text((x, y), "\n".join(lines), font=fnt, fill=fill, spacing=6, align="center")


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#3F6382", width: int = 4):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=fill, width=width)
    if abs(x2 - x1) >= abs(y2 - y1):
        if x2 >= x1:
            pts = [(x2, y2), (x2 - 18, y2 - 9), (x2 - 18, y2 + 9)]
        else:
            pts = [(x2, y2), (x2 + 18, y2 - 9), (x2 + 18, y2 + 9)]
    else:
        if y2 >= y1:
            pts = [(x2, y2), (x2 - 9, y2 - 18), (x2 + 9, y2 - 18)]
        else:
            pts = [(x2, y2), (x2 - 9, y2 + 18), (x2 + 9, y2 + 18)]
    draw.polygon(pts, fill=fill)


def box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, body: Iterable[str], fill: str, outline: str):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=4)
    title_font = font(26)
    body_font = font(20)
    draw.text((xy[0] + 26, xy[1] + 20), title, font=title_font, fill="#173B57")
    y = xy[1] + 66
    for item in body:
        for line in wrap_text(draw, item, body_font, xy[2] - xy[0] - 52):
            draw.text((xy[0] + 28, y), line, font=body_font, fill="#344B5E")
            y += 28
        y += 4


def create_architecture_diagram(path: Path):
    img = Image.new("RGB", (1700, 1100), "white")
    draw = ImageDraw.Draw(img)
    title = font(38)
    draw.text((70, 42), "Jasmine 软件体系结构图", font=title, fill="#103A5C")

    box(draw, (80, 135, 520, 380), "用户界面层",
        ["App 负责视图路由与响应式布局", "Sidebar、DiscoverView、SearchView、LibraryView、NowPlayingView 等组件承载交互", "PlayerBar 提供全局播放控制"],
        "#F4F8FB", "#7EA2C0")
    box(draw, (630, 135, 1070, 380), "业务状态层",
        ["PlayerContext 统一管理歌曲、歌单、队列、播放状态", "useKeyboardShortcuts 处理桌面快捷键", "i18n 上下文提供中英文界面"],
        "#F7FAF4", "#84A971")
    box(draw, (1180, 135, 1620, 380), "服务适配层",
        ["neteaseApi 封装搜索、登录、歌曲地址、歌词、排行榜", "MusicSource 抽象预留多音源扩展", "analyticsApi 记录播放行为"],
        "#FBF7F0", "#BE9C63")

    box(draw, (210, 560, 650, 830), "本地持久化",
        ["IndexedDB 保存音频二进制文件", "localStorage 保存歌曲元数据、歌单、历史、登录 Cookie 与播放状态", "music-metadata 解析本地音频标签"],
        "#F8F5FA", "#9E83B8")
    box(draw, (760, 560, 1200, 830), "桌面运行层",
        ["Electron 主进程创建窗口并启动本地服务", "preload 暴露受控桌面能力", "HTML5 Audio 执行播放"],
        "#F4F8FB", "#7EA2C0")
    box(draw, (1310, 560, 1620, 830), "外部服务",
        ["NeteaseCloudMusicApi", "网易云音乐接口", "本地分析服务 3001"],
        "#FFF7F7", "#C98282")

    arrow(draw, (520, 255), (630, 255))
    arrow(draw, (1070, 255), (1180, 255))
    arrow(draw, (850, 380), (430, 560))
    arrow(draw, (850, 380), (980, 560))
    arrow(draw, (1400, 380), (1470, 560))
    arrow(draw, (1180, 255), (1120, 560))
    centered_text(draw, (100, 920, 1600, 1030),
                  "设计要点：界面组件只负责展示与触发操作；播放器上下文集中处理业务状态；服务层隔离在线接口；持久化层分别处理大文件与结构化元数据。",
                  font(24), "#334E63")
    img.save(path)


def create_ui_flow_diagram(path: Path):
    img = Image.new("RGB", (1700, 1060), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 42), "Jasmine 用户界面与跳转关系", font=font(38), fill="#103A5C")

    nav = (90, 150, 430, 840)
    draw.rounded_rectangle(nav, radius=24, fill="#17324A", outline="#5A7893", width=4)
    draw.text((130, 190), "侧边栏导航", font=font(30), fill="white")
    nav_items = ["发现 discover", "音乐库 library", "搜索 search", "下载 download", "设置 settings", "歌单 playlist"]
    centers = []
    for i, text in enumerate(nav_items):
        y = 260 + i * 82
        draw.rounded_rectangle((125, y, 395, y + 52), radius=12, fill="#264B6C", outline="#8FB0CA", width=2)
        centered_text(draw, (125, y, 395, y + 52), text, font(22), "#EAF4FC")
        centers.append((395, y + 26))

    screens = [
        ("发现页", "热门搜索 / 推荐歌单 / 排行榜 / 艺人页", (610, 150, 1030, 270)),
        ("搜索页", "本地与在线结果检索，进入歌曲或歌单", (1180, 150, 1600, 270)),
        ("音乐库", "本地歌曲 / 自建歌单 / 收藏专辑 / 历史", (610, 360, 1030, 480)),
        ("歌单详情", "播放歌单、添加歌曲、封面与名称管理", (1180, 360, 1600, 480)),
        ("全屏播放", "歌词、歌曲信息、相似歌曲、分享卡片", (610, 570, 1030, 690)),
        ("播放队列", "下一首播放、移除、清空队列", (1180, 570, 1600, 690)),
        ("设置页", "语言切换，后续扩展系统配置", (895, 780, 1315, 900)),
    ]
    for name, desc, xy in screens:
        draw.rounded_rectangle(xy, radius=18, fill="#F4F8FB", outline="#7EA2C0", width=4)
        draw.text((xy[0] + 24, xy[1] + 18), name, font=font(26), fill="#173B57")
        centered_text(draw, (xy[0] + 24, xy[1] + 56, xy[2] - 24, xy[3] - 15), desc, font(20), "#344B5E")

    for start, target in zip(centers[:5], [screens[0], screens[2], screens[1], screens[5], screens[6]]):
        xy = target[2]
        arrow(draw, start, (xy[0], (xy[1] + xy[3]) // 2), "#7795AD", 3)
    arrow(draw, (1030, 420), (1180, 420), "#7795AD", 3)
    arrow(draw, (820, 690), (820, 780), "#7795AD", 3)
    arrow(draw, (1030, 630), (1180, 630), "#7795AD", 3)

    draw.rounded_rectangle((90, 930, 1610, 1010), radius=16, fill="#EEF3F7", outline="#BACBDD", width=2)
    centered_text(draw, (120, 938, 1580, 1002), "底部 PlayerBar 在普通视图中常驻，负责当前歌曲展示、进度控制、音量、播放模式与队列入口；NowPlaying 视图隐藏侧边栏以获得沉浸式播放体验。", font(23), "#31495E")
    img.save(path)


def create_sequence_diagram(path: Path):
    img = Image.new("RGB", (1760, 1120), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 42), "用例实现时序图：在线搜索并播放歌曲", font=font(38), fill="#103A5C")
    participants = [
        ("用户", 140),
        ("Discover/SearchView", 460),
        ("neteaseApi", 790),
        ("PlayerContext", 1120),
        ("HTML5 Audio", 1450),
    ]
    top, bottom = 150, 1030
    for name, x in participants:
        draw.rounded_rectangle((x - 110, top, x + 110, top + 60), radius=14, fill="#F4F8FB", outline="#7EA2C0", width=3)
        centered_text(draw, (x - 106, top + 8, x + 106, top + 52), name, font(22), "#173B57")
        draw.line((x, top + 60, x, bottom), fill="#B8C8D6", width=2)

    steps = [
        (140, 460, 270, "输入关键词或点击推荐内容"),
        (460, 790, 360, "searchOnline / getPlaylistDetail"),
        (790, 460, 450, "返回 Song、Playlist 数据"),
        (140, 460, 540, "点击歌曲播放"),
        (460, 1120, 630, "playSong(song, context)"),
        (1120, 790, 720, "按需请求 song/url 与歌词"),
        (790, 1120, 810, "返回 audioUrl 与 LyricLine[]"),
        (1120, 1450, 900, "设置 src、load、play"),
        (1450, 1120, 990, "timeupdate / duration / ended 事件"),
    ]
    for x1, x2, y, label in steps:
        arrow(draw, (x1, y), (x2, y), "#3F6382", 3)
        draw.multiline_text((min(x1, x2) + 20, y - 42), "\n".join(wrap_text(draw, label, font(18), abs(x2 - x1) - 42)), font=font(18), fill="#344B5E", spacing=4)
    img.save(path)


def create_class_diagram(path: Path):
    img = Image.new("RGB", (1800, 1340), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 42), "Jasmine 核心实现类图（TypeScript 接口与模块）", font=font(38), fill="#103A5C")

    def class_box(x: int, y: int, w: int, h: int, name: str, attrs: list[str], methods: list[str] | None = None):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=12, fill="#F9FBFD", outline="#6F93AF", width=3)
        draw.rectangle((x, y, x + w, y + 50), fill="#E8F0F7", outline="#6F93AF", width=3)
        centered_text(draw, (x, y + 4, x + w, y + 46), name, font(22), "#173B57")
        small = font(17)
        yy = y + 66
        for attr in attrs:
            draw.text((x + 18, yy), attr, font=small, fill="#344B5E")
            yy += 25
        if methods:
            draw.line((x, yy + 8, x + w, yy + 8), fill="#AFC1D0", width=2)
            yy += 22
            for method in methods:
                draw.text((x + 18, yy), method, font=small, fill="#344B5E")
                yy += 25

    class_box(70, 140, 360, 250, "Song",
              ["id: string", "title / artist / album", "duration: number", "source: local | online", "neteaseId? / audioUrl?", "lyrics?: LyricLine[]"])
    class_box(520, 140, 360, 230, "Playlist",
              ["id / name / description", "coverColor: string", "songs: Song[]", "createdAt: number", "creator?: string"])
    class_box(970, 140, 360, 230, "PlayerState",
              ["currentSong: Song | null", "isPlaying / currentTime", "volume / repeatMode", "queue: Song[]", "queueIndex: number"])
    class_box(1420, 140, 300, 180, "LyricLine",
              ["time: number", "text: string"])

    class_box(70, 480, 520, 360, "PlayerContext",
              ["state: PlayerState", "userSongs: Song[]", "userPlaylists: Playlist[]", "savedAlbums: SavedAlbum[]", "playHistory: PlayRecord[]"],
              ["playSong(song, context?)", "playPlaylist(playlist)", "importFiles(files)", "create/rename/deletePlaylist", "addToQueue / playNext", "toggleLike / toggleAlbum"])
    class_box(700, 480, 430, 340, "MusicSource",
              ["id / name / color", "面向在线音源的统一接口"],
              ["search(keywords)", "getPlaylistDetail(id)", "getAudioUrl(id)", "getLyrics(id)", "getRecommendations()", "getCharts()"])
    class_box(1240, 480, 430, 340, "neteaseSource / neteaseApi",
              ["网易云音源适配器", "Cookie 登录状态", "音频地址缓存"],
              ["searchOnline()", "getPlaylistDetail()", "getSongAudioUrl()", "getLyrics()", "getUserPlaylists()"])

    class_box(170, 910, 430, 320, "db.ts",
              ["IndexedDB: audioFiles", "localStorage: mp-songs", "mp-playlists / mp-albums", "mp-play-history"],
              ["saveAudioFile()", "loadAudioFile()", "saveSongs()", "savePlaylists()"])
    class_box(720, 910, 410, 270, "UI Components",
              ["Sidebar / PlayerBar", "DiscoverView / SearchView", "LibraryView / PlaylistDetail", "NowPlayingView / QueueView"],
              ["触发上下文操作", "展示当前状态"])
    class_box(1240, 910, 430, 320, "Electron Main",
              ["BrowserWindow", "Netease API server: 3000", "Analytics server: 3001", "download-audio IPC"],
              ["createWindow()", "serveNcmApi()", "startAnalyticsServer()"])

    arrow(draw, (430, 260), (520, 260), "#7795AD", 3)
    arrow(draw, (970, 260), (880, 260), "#7795AD", 3)
    arrow(draw, (1330, 230), (1420, 230), "#7795AD", 3)
    arrow(draw, (330, 480), (250, 390), "#7795AD", 3)
    arrow(draw, (330, 480), (760, 370), "#7795AD", 3)
    arrow(draw, (330, 840), (385, 910), "#7795AD", 3)
    arrow(draw, (590, 650), (700, 650), "#7795AD", 3)
    arrow(draw, (1130, 650), (1240, 650), "#7795AD", 3)
    arrow(draw, (925, 910), (330, 840), "#7795AD", 3)
    arrow(draw, (1455, 910), (1455, 820), "#7795AD", 3)
    img.save(path)


def create_data_deployment_diagram(path: Path):
    img = Image.new("RGB", (1750, 1180), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 42), "数据与部署设计图", font=font(38), fill="#103A5C")

    box(draw, (90, 140, 520, 390), "用户桌面环境",
        ["Windows / macOS", "Jasmine Electron 应用", "本地音频文件导入"],
        "#F4F8FB", "#7EA2C0")
    box(draw, (660, 140, 1090, 390), "渲染进程",
        ["React + TypeScript UI", "PlayerContext 状态中心", "HTML5 Audio 播放引擎"],
        "#F7FAF4", "#84A971")
    box(draw, (1230, 140, 1660, 390), "主进程与本地服务",
        ["Electron main.cjs", "Netease API: 127.0.0.1:3000", "Analytics API: 127.0.0.1:3001"],
        "#FBF7F0", "#BE9C63")
    box(draw, (210, 590, 640, 870), "浏览器本地存储",
        ["IndexedDB/audioFiles: 音频二进制", "localStorage/mp-songs: 歌曲元数据", "mp-playlists: 歌单及歌曲引用", "mp-albums、mp-play-history、mp-playback-state"],
        "#F8F5FA", "#9E83B8")
    box(draw, (790, 590, 1220, 870), "应用数据文件",
        ["analytics.sqlite", "打包后的 dist 静态资源", "Electron release 安装包"],
        "#F4F8FB", "#7EA2C0")
    box(draw, (1370, 590, 1660, 870), "外部网络",
        ["网易云音乐接口", "歌单 / 歌曲 / 歌词 / 登录", "在线音频 URL"],
        "#FFF7F7", "#C98282")

    arrow(draw, (520, 265), (660, 265))
    arrow(draw, (1090, 265), (1230, 265))
    arrow(draw, (875, 390), (425, 590))
    arrow(draw, (1445, 390), (1445, 590))
    arrow(draw, (1090, 735), (790, 735))
    arrow(draw, (1230, 265), (1515, 590))
    centered_text(draw, (120, 970, 1630, 1080),
                  "部署关系：桌面应用启动后由主进程创建窗口并启动本地服务；渲染进程访问本地服务和浏览器存储；网易云接口作为外部依赖，音频文件与业务元数据保存在用户本机。",
                  font(24), "#334E63")
    img.save(path)


def set_run_font(run, name: str = "宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    set_run_font(run)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def format_table(table, header_fill: str = "E8EEF5"):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B8C8D6")
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run)
                    run.font.size = Pt(10.5)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_heading(doc: Document, text: str, level: int, page_break_before: bool = False):
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    p.paragraph_format.page_break_before = page_break_before
    if level == 1:
        p.paragraph_format.keep_with_next = True
    if level in (2, 3):
        p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, "黑体")
    return p


def add_para(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        set_run_font(r1)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(4)
    marker = p.add_run("• ")
    set_run_font(marker)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_image(doc: Document, path: Path, caption: str, width: float = 6.2):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    r = c.add_run(caption)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(80, 80, 80)
    set_run_font(r)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def update_cover(doc: Document):
    for p in doc.paragraphs:
        text = p.text.strip()
        replacement = None
        if text == "文档编号：<项目名称> – SDS – <**.**>":
            replacement = "文档编号：Jasmine - SDS - V1.0"
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text == "<项目名称>":
            replacement = "Jasmine"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
        elif text == "软件设计规格说明书":
            replacement = "软件设计规格说明书"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(18)
        elif text == "日期：":
            replacement = f"日期：{DATE_TEXT}"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text == "文档变更历史记录":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(8)

        if replacement is not None:
            p.text = replacement
        if replacement is not None or text == "文档变更历史记录":
            for run in p.runs:
                set_run_font(run, "黑体" if text in {"<项目名称>", "软件设计规格说明书", "文档变更历史记录"} else "宋体")
                if text == "<项目名称>":
                    run.font.size = Pt(28)
                    run.bold = True
                    run.font.color.rgb = RGBColor(16, 58, 92)
                elif text == "软件设计规格说明书":
                    run.font.size = Pt(22)
                    run.bold = True
                    run.font.color.rgb = RGBColor(31, 77, 120)
                elif text == "文档变更历史记录":
                    run.font.size = Pt(14)
                    run.bold = True
                    run.font.color.rgb = RGBColor(31, 77, 120)
                else:
                    run.font.size = Pt(10.5)


def rebuild_front_matter(doc: Document):
    date_idx = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("日期：")), None)
    body_idx = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "1、引言"), None)
    if date_idx is None or body_idx is None:
        return

    date_para = doc.paragraphs[date_idx]
    body_para = doc.paragraphs[body_idx]
    body_el = doc._body._element
    children = list(body_el)
    start = children.index(date_para._element) + 1
    end = children.index(body_para._element)
    for child in children[start:end]:
        body_el.remove(child)

    toc_entries = [
        "目录",
        "1、引言",
        "1.1 编写目的",
        "1.2 读者对象",
        "1.3 软件项目概述",
        "1.4 文档概述",
        "1.5 定义",
        "1.6 参考资料",
        "2、软件设计约束",
        "2.1 软件设计目标和原则",
        "2.2 软件设计的约束和限制",
        "2.3 框架选择与技术路线",
        "3、软件设计",
        "3.1 软件体系结构设计",
        "3.2 用户界面设计",
        "3.3 用例设计",
        "3.4 类设计",
        "3.5 数据设计",
        "3.6 部署设计",
    ]

    toc_title_para = None
    for text in toc_entries:
        p = body_para.insert_paragraph_before(text)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if text == "目录":
            toc_title_para = p
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.page_break_before = True
            p.paragraph_format.space_after = Pt(14)
        elif text[0].isdigit() and "、" in text:
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.space_before = Pt(4)
        elif text.count(".") == 1:
            p.paragraph_format.left_indent = Inches(0.3)
        elif text.count(".") >= 2:
            p.paragraph_format.left_indent = Inches(0.55)
        for run in p.runs:
            set_run_font(run, "黑体" if text == "目录" else "宋体")
            if text == "目录":
                run.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(31, 77, 120)
            elif text[0].isdigit() and "、" in text:
                run.bold = True

    if toc_title_para is None:
        return

    history_heading = toc_title_para.insert_paragraph_before("文档变更历史记录")
    history_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    history_heading.paragraph_format.page_break_before = True
    history_heading.paragraph_format.space_before = Pt(10)
    history_heading.paragraph_format.space_after = Pt(10)
    heading_run = history_heading.runs[0]
    set_run_font(heading_run, "黑体")
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    heading_run.font.color.rgb = RGBColor(31, 77, 120)

    history_table = doc.add_table(rows=2, cols=5)
    headers = ["序号", "变更日期", "变更人员", "变更内容详情描述", "版本"]
    values = ["1", DATE_TEXT, "项目组", "依据 Jasmine 项目源码和说明文档完成首版软件设计规格说明书", "V1.0"]
    widths = [0.6, 1.2, 1.0, 3.4, 0.7]
    for i, text in enumerate(headers):
        set_cell_text(history_table.rows[0].cells[i], text, True)
        history_table.rows[0].cells[i].width = Inches(widths[i])
    for i, text in enumerate(values):
        set_cell_text(history_table.rows[1].cells[i], text)
        history_table.rows[1].cells[i].width = Inches(widths[i])
    format_table(history_table)
    toc_title_para._p.addprevious(history_table._tbl)

    spacer = toc_title_para.insert_paragraph_before("")
    spacer.paragraph_format.space_after = Pt(6)


def clear_body(doc: Document):
    body_indices = [i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "1、引言"]
    body_start = body_indices[-1] if body_indices else None
    if body_start is None:
        return
    for p in list(doc.paragraphs[body_start:])[::-1]:
        delete_paragraph(p)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    for paragraph in doc.paragraphs[-1:]:
        paragraph.paragraph_format.space_after = Pt(4)
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            set_cell_text(row.cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    format_table(table)
    doc.add_paragraph()
    return table


def build_content(doc: Document):
    add_heading(doc, "1、引言", 1, page_break_before=True)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(doc, "本文档用于说明 Jasmine 桌面音乐播放器的软件设计方案，重点描述系统在框架选择、体系结构、界面组织、用例实现、核心类/接口、数据持久化和部署运行方面的设计。文档以当前项目源码、README、需求分析材料和构建配置为依据，为课程期末作业、毕业设计答辩、后续维护和测试验收提供统一的设计参考。")

    add_heading(doc, "1.2 读者对象", 2)
    add_para(doc, "本文档面向项目开发者、课程教师、评审人员、测试人员以及后续可能接手维护 Jasmine 的协作者。开发者可据此理解模块边界和关键接口，评审人员可据此判断系统设计的完整性和合理性，测试人员可据此设计覆盖核心链路的测试用例。")

    add_heading(doc, "1.3 软件项目概述", 2)
    add_para(doc, "项目名称为 Jasmine，项目代号为 MusicPlayer--Jasmine。该项目是一个基于 Vite、React 18、TypeScript 与 Electron 构建的桌面音乐播放器，支持本地音乐导入、网易云在线搜索与播放、歌单管理、歌词同步、播放历史、队列管理、中英文国际化、下载和桌面端打包发布。")
    add_para(doc, "系统的核心目标是在一个统一桌面应用中整合本地音乐管理与在线音乐发现两类场景，使用户能够完成“导入/搜索歌曲、组织歌单、播放控制、查看歌词、恢复播放状态、桌面端持续使用”的完整闭环。")

    add_heading(doc, "1.4 文档概述", 2)
    add_para(doc, "本文档共三大部分：第一部分说明文档目标、读者、项目概况和术语；第二部分说明设计目标、设计原则、运行约束和技术框架选择；第三部分从体系结构、用户界面、用例、类设计、数据设计和部署设计六个方面展开软件设计。")

    add_heading(doc, "1.5 定义", 2)
    add_table(doc, ["术语", "含义"], [
        ["SDS", "Software Design Specification，软件设计规格说明书。"],
        ["Electron 主进程", "负责创建桌面窗口、启动本地服务并处理下载 IPC 的 Node.js 运行环境。"],
        ["渲染进程", "运行 React 前端界面、播放状态和用户交互逻辑的浏览器环境。"],
        ["PlayerContext", "项目中的播放器业务状态中心，统一维护当前歌曲、队列、音量、循环模式、歌单和历史记录等数据。"],
        ["MusicSource", "在线音源抽象接口，用于隔离具体音乐平台的搜索、歌单、音频地址、歌词和推荐能力。"],
        ["IndexedDB", "浏览器本地数据库，本项目用于保存导入音频文件的二进制数据。"],
        ["localStorage", "浏览器键值存储，本项目用于保存歌曲元数据、歌单、收藏、历史、Cookie 和播放状态。"],
    ], [1.5, 5.0])

    add_heading(doc, "1.6 参考资料", 2)
    for item in [
        "Jasmine README.md：项目功能、技术栈、运行方式与目录结构说明。",
        "package.json：依赖库、构建脚本、Electron Builder 发布配置。",
        "src/context/PlayerContext.tsx：播放器状态管理和核心业务方法实现。",
        "src/lib/db.ts：IndexedDB 与 localStorage 持久化实现。",
        "src/lib/neteaseApi.ts 与 src/lib/sources/*：网易云 API 封装与音源抽象设计。",
        "electron/main.cjs：桌面主进程、本地 API 服务、下载 IPC 与分析服务启动逻辑。",
        "docs/jasmine-requirements-analysis.md：Jasmine 项目需求分析材料。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2、软件设计约束", 1)
    add_heading(doc, "2.1 软件设计目标和原则", 2)
    add_para(doc, "Jasmine 的设计目标是构建一个功能闭环完整、交互体验一致、可桌面化发布、具备后续扩展空间的音乐播放器系统。系统既要能完成本地歌曲导入与长期管理，也要能接入网易云在线内容，并以统一播放队列和统一底部播放器串联不同页面产生的播放行为。")
    add_para(doc, "设计原则包括：一是分层解耦，界面层不直接承担复杂播放与存储逻辑；二是状态集中，播放相关状态由 PlayerContext 管理；三是接口隔离，在线音乐能力通过 neteaseApi 与 MusicSource 抽象封装；四是本地优先，用户导入的音频与业务元数据保存在本机；五是渐进扩展，当前稳定接入网易云，同时保留接入其他音源和分析能力的空间。")

    add_heading(doc, "2.2 软件设计的约束和限制", 2)
    add_table(doc, ["约束类别", "具体约束"], [
        ["运行环境", "桌面端以 Electron 运行为主，开发环境可通过 Vite 浏览器地址访问；目标平台主要覆盖 macOS 与 Windows。"],
        ["开发语言", "前端和业务逻辑采用 TypeScript/React；Electron 主进程与本地服务脚本采用 CommonJS JavaScript。"],
        ["在线能力", "网易云搜索、歌单、歌词、登录和在线播放依赖本地 NeteaseCloudMusicApi 服务，默认端口为 3000。"],
        ["本地存储", "音频文件保存在 IndexedDB，结构化业务数据保存在 localStorage；因此大规模音频数据受浏览器存储配额影响。"],
        ["安全边界", "Electron 配置禁用 nodeIntegration 并启用 contextIsolation，渲染进程通过 preload 暴露的受控 API 与主进程交互。"],
        ["性能边界", "大歌单、批量音频地址请求和在线歌词获取存在网络延迟，需要通过批量请求、缓存和渐进加载降低等待成本。"],
        ["课程交付", "文档和实现需突出结构清晰、功能完整、可演示、可说明，而不是引入过度复杂的企业级基础设施。"],
    ], [1.4, 5.1])

    add_heading(doc, "2.3 框架选择与技术路线", 2)
    add_para(doc, "项目选择 Vite + React + TypeScript + Electron 的组合，原因是该组合能够同时满足快速前端开发、类型安全、组件化界面组织和桌面端发布需要。对于音乐播放器而言，React 适合组织复杂交互界面，TypeScript 适合约束歌曲、歌单、队列、音源等业务对象，Electron 则将 Web 能力扩展为可安装桌面软件。")
    add_table(doc, ["框架/技术", "在项目中的用途", "选择理由"], [
        ["Vite 6", "前端构建与开发服务器", "启动快、配置简洁，适合 React 单页应用开发和桌面端构建。"],
        ["React 18", "界面组件与视图组织", "组件化表达清晰，适合发现页、音乐库、播放页、队列等复杂视图组合。"],
        ["TypeScript", "业务类型与接口约束", "通过 Song、Playlist、PlayerState、MusicSource 等类型降低状态流转错误。"],
        ["Electron 41", "桌面应用封装", "提供安装包、窗口管理、主进程下载能力和本地服务启动能力。"],
        ["IndexedDB", "保存导入音频文件", "适合在浏览器环境中保存较大的二进制数据。"],
        ["localStorage", "保存元数据和偏好", "实现简单，适合歌曲元数据、歌单、历史、播放状态等结构化小数据。"],
        ["NeteaseCloudMusicApi", "网易云在线音乐接入", "为搜索、歌单、歌曲地址、歌词、登录提供可用接口基础。"],
        ["music-metadata", "本地音频元数据解析", "用于从导入文件中读取歌曲名、歌手、专辑、时长等信息。"],
    ], [1.35, 2.1, 3.05])

    add_heading(doc, "3、软件设计", 1)
    add_heading(doc, "3.1 软件体系结构设计", 2)
    add_para(doc, "Jasmine 采用分层式桌面应用结构。最外层是 React 组件组成的用户界面层；中间是 PlayerContext、Hook、i18n 和在线服务封装组成的业务状态与服务层；底层是 IndexedDB/localStorage、本地 API 服务、Electron 主进程和外部网易云接口组成的运行支撑层。")
    add_image(doc, ASSET_DIR / "architecture.png", "图 3-1 Jasmine 软件体系结构图")
    add_para(doc, "在逻辑模型上，App.tsx 负责主布局、视图路由、侧边栏收起和移动端导航；components 目录承载具体页面；PlayerContext 负责播放状态、歌单、收藏、历史和导入下载等业务能力；lib 目录负责持久化、网易云 API、歌词解析、元数据解析和音源抽象；electron 目录负责桌面窗口、本地服务和下载能力。")

    add_table(doc, ["层次", "主要文件/模块", "设计职责"], [
        ["界面层", "App.tsx、components/*、App.css", "组织导航、主内容视图、底部播放器、歌词页、歌单详情和响应式布局。"],
        ["状态层", "PlayerContext.tsx、useKeyboardShortcuts.ts", "管理播放状态、队列、歌单、收藏、历史、导入、快捷键与恢复逻辑。"],
        ["服务层", "neteaseApi.ts、sources/*、lyrics.ts、metadata.ts", "封装在线接口、音源抽象、歌词解析和本地文件元数据读取。"],
        ["持久化层", "db.ts、localStorage、IndexedDB", "保存音频文件、歌曲元数据、歌单、专辑收藏、历史记录和播放状态。"],
        ["桌面层", "electron/main.cjs、preload.cjs、analytics/server.cjs", "创建窗口、启动本地网易云 API 与分析服务、处理跨进程下载。"],
    ], [1.1, 2.0, 3.4])

    add_heading(doc, "3.2 用户界面设计", 2)
    add_para(doc, "界面采用“侧边栏导航 + 主内容区域 + 底部播放器”的桌面音乐软件布局。侧边栏提供发现、音乐库、搜索、下载、设置和歌单入口；主内容区域根据当前视图渲染 DiscoverView、SearchView、LibraryView、PlaylistDetail、NowPlayingView、QueueView 等页面；底部 PlayerBar 在普通视图中持续显示当前歌曲和播放控制。")
    add_image(doc, ASSET_DIR / "ui_flow.png", "图 3-2 Jasmine 用户界面与跳转关系")
    add_para(doc, "界面设计重点是保证播放控制始终可达、页面职责清晰、在线内容和本地内容在导航上自然合并。NowPlayingView 作为沉浸式播放页，会隐藏普通侧边栏并展示歌词、歌曲详情、相似歌曲和分享入口；QueueView 作为临时管理页面，从底部播放器进入并可返回原视图。")
    add_table(doc, ["界面", "主要职责", "对应组件"], [
        ["发现页", "展示热门搜索、推荐歌单、排行榜、艺人和专辑内容。", "DiscoverView"],
        ["音乐库", "展示本地歌曲、自建歌单、网易云用户歌单、收藏专辑和历史入口。", "LibraryView"],
        ["搜索页", "完成歌曲、歌单、艺人、专辑等内容检索并进入播放或详情。", "SearchView"],
        ["歌单详情", "展示歌曲列表，支持播放歌单、添加/移除歌曲、修改封面。", "PlaylistDetail"],
        ["全屏播放页", "展示歌词同步、当前歌曲信息、相似歌曲和分享卡片。", "NowPlayingView"],
        ["播放队列", "展示当前队列并支持下一首、移除和清空操作。", "QueueView"],
        ["设置页", "提供语言切换，并作为后续主题、缓存、快捷键等设置入口。", "SettingsView"],
    ], [1.2, 3.0, 2.3])

    add_heading(doc, "3.3 用例设计", 2)
    add_para(doc, "系统关键用例包括本地音乐导入、在线搜索并播放、歌单管理、歌词查看、播放队列管理、网易云登录和桌面端持续使用。其中在线搜索并播放是连接界面层、服务层、播放器状态层和音频播放引擎的典型用例。")
    add_image(doc, ASSET_DIR / "sequence_online_play.png", "图 3-3 在线搜索并播放歌曲时序图")
    add_para(doc, "在线搜索并播放的基本流程为：用户在发现页或搜索页输入关键词；界面调用 neteaseApi 或 MusicSource 获取在线结果；用户点击歌曲后，界面调用 PlayerContext.playSong；PlayerContext 根据 Song 中的 neteaseId 按需获取音频地址和歌词，并设置 HTML5 Audio 的 src；播放事件再反向同步进度、时长和结束后的下一首调度。")
    add_table(doc, ["用例编号", "用例名称", "参与者", "主要设计对象", "结果"], [
        ["UC-01", "本地音乐导入", "用户、本地文件系统", "LibraryView、PlayerContext.importFiles、metadata、db.ts", "歌曲文件进入 IndexedDB，歌曲元数据进入 localStorage，并出现在音乐库。"],
        ["UC-02", "在线搜索并播放", "用户、网易云服务", "Discover/SearchView、neteaseApi、PlayerContext、HTML5 Audio", "在线歌曲加入队列并开始播放，歌词和进度同步。"],
        ["UC-03", "歌单管理", "用户", "CreatePlaylistModal、PlaylistDetail、PlayerContext", "歌单被创建、重命名、删除、添加歌曲或更新封面，并持久化保存。"],
        ["UC-04", "播放队列管理", "用户", "QueueView、PlayerBar、PlayerContext reducer", "歌曲可添加到队列、下一首播放、移除或清空，当前播放上下文保持一致。"],
        ["UC-05", "桌面下载在线歌曲", "用户、Electron 主进程", "DownloadView、preload、ipcMain download-audio、db.ts", "主进程完成网络下载并将音频数据回传保存到本地。"],
    ], [0.9, 1.4, 1.2, 2.4, 2.5])

    add_heading(doc, "3.4 类设计", 2)
    add_para(doc, "项目主体使用 TypeScript 接口、React Context 和函数模块表达领域模型。核心模型包括 Song、Playlist、SavedAlbum、PlayerState、LyricLine 和 MusicSource；核心控制对象是 PlayerContext；核心适配对象是 neteaseApi 与 neteaseSource；核心持久化对象集中在 db.ts。")
    add_image(doc, ASSET_DIR / "class_diagram.png", "图 3-4 Jasmine 核心实现类图")
    add_table(doc, ["类/接口/模块", "可见范围", "关键属性/方法", "说明"], [
        ["Song", "src/types.ts 导出", "id、title、artist、album、duration、source、neteaseId、audioUrl、lyrics", "歌曲领域对象，兼容本地歌曲和在线歌曲。"],
        ["Playlist", "src/types.ts 导出", "id、name、description、coverColor、songs、createdAt、creator", "歌单领域对象，既可表示自建歌单，也可映射网易云歌单。"],
        ["PlayerState", "src/types.ts 导出", "currentSong、isPlaying、currentTime、volume、repeatMode、queue、queueIndex", "播放状态快照，由 reducer 维护。"],
        ["MusicSource", "src/lib/sources/types.ts 导出", "search、getPlaylistDetail、getAudioUrl、getBatchAudioUrls、getLyrics、getRecommendations", "在线音源统一接口，为后续多音源接入保留扩展点。"],
        ["PlayerContext", "React Context", "playSong、playPlaylist、importFiles、createPlaylist、addToQueue、toggleLike、downloadOnlineSong", "系统业务中枢，向界面提供状态和操作方法。"],
        ["db.ts", "lib 内部模块导出", "saveAudioFile、loadAudioFile、saveSongs、savePlaylists、savePlayHistory", "封装本地数据保存和读取。"],
        ["neteaseApi.ts", "lib 内部模块导出", "searchOnline、getPlaylistDetail、getSongAudioUrl、getLyrics、getLoginStatus", "封装网易云在线接口和缓存逻辑。"],
    ], [1.35, 1.0, 2.35, 1.75])

    add_heading(doc, "3.5 数据设计", 2)
    add_para(doc, "Jasmine 没有采用传统服务端关系数据库作为主存储，而是基于桌面 Web 容器的本地存储能力进行数据设计。音频二进制体积较大，保存在 IndexedDB 的 audioFiles 对象仓库；歌曲、歌单、收藏、历史和播放状态数据较小，使用 localStorage 的 JSON 字符串保存；本地分析服务额外使用 analytics.sqlite 保存播放行为分析数据。")
    add_table(doc, ["数据项", "存储位置", "键/对象仓库", "主要字段", "设计说明"], [
        ["音频文件", "IndexedDB", "music-player-store / audioFiles", "id、data、type", "保存用户导入或下载的音频二进制数据，并生成 Object URL 播放。"],
        ["歌曲元数据", "localStorage", "mp-songs", "id、title、artist、album、duration、coverUrl、source、neteaseId、lyrics", "保存歌曲结构化信息，避免重启后丢失音乐库。"],
        ["歌单数据", "localStorage", "mp-playlists", "id、name、description、coverColor、songs、_embeddedSongs", "songs 保存歌曲 id 引用，同时嵌入歌曲快照提高恢复能力。"],
        ["收藏专辑", "localStorage", "mp-albums", "id、neteaseId、name、artist、picUrl、savedAt", "保存用户收藏的在线专辑。"],
        ["播放历史", "localStorage", "mp-play-history", "songId、playedAt", "记录最近播放行为，当前最多保留 500 条。"],
        ["播放状态", "localStorage", "mp-playback-state", "currentSong、currentTime、volume、repeatMode、queue、queueIndex、savedAt", "用于应用重启后恢复当前歌曲、队列和音量等状态。"],
        ["登录 Cookie", "localStorage", "mp-netease-cookie", "cookie 字符串", "保存网易云登录状态，用于恢复用户歌单和更完整的在线能力。"],
        ["分析数据", "SQLite", "tmp/analytics.sqlite 或 userData/analytics.sqlite", "播放事件记录", "供本地分析服务统计播放行为。"],
    ], [1.0, 1.15, 1.35, 2.15, 2.0])
    add_image(doc, ASSET_DIR / "data_deployment.png", "图 3-5 数据与部署设计图")
    add_para(doc, "数据恢复流程由 PlayerProvider 初始化阶段完成：先读取 mp-songs 和 IndexedDB 中的音频文件，再根据 mp-playlists 中保存的歌曲 id 或 _embeddedSongs 还原歌单，随后加载收藏专辑、播放历史和播放状态。该设计降低了单一存储项损坏导致全部业务数据不可恢复的风险。")

    add_heading(doc, "3.6 部署设计", 2)
    add_para(doc, "系统部署形态分为开发运行和桌面发布两种。开发运行时，通过 npm run dev 同时启动 Vite 前端、本地网易云 API 服务和本地分析服务；桌面发布时，通过 Electron Builder 将 dist 静态资源和 electron 目录打包为 macOS DMG/ZIP 或 Windows NSIS 安装包。")
    add_table(doc, ["部署节点", "部署内容", "端口/路径", "说明"], [
        ["用户桌面操作系统", "Jasmine 安装包或源码运行环境", "macOS / Windows", "承载应用窗口、音频播放和本地存储。"],
        ["Electron 主进程", "main.cjs、preload.cjs", "应用内部", "创建 BrowserWindow，启动本地服务，处理下载 IPC。"],
        ["React 渲染进程", "dist/index.html 与打包资源", "开发端口 5173 或 file://dist", "展示界面并通过 PlayerContext 管理业务状态。"],
        ["网易云本地 API", "NeteaseCloudMusicApi", "127.0.0.1:3000", "提供搜索、歌单、歌词、登录和音频地址接口。"],
        ["分析服务", "analytics/server.cjs", "127.0.0.1:3001", "记录播放行为并写入 SQLite。"],
        ["外部网易云服务", "网易云音乐数据源", "HTTPS 网络访问", "为本地 API 提供最终在线数据来源。"],
    ], [1.2, 1.6, 1.2, 2.5])
    add_para(doc, "部署设计的关键点是将在线 API 服务与桌面窗口生命周期绑定。应用 ready 后尝试启动 NeteaseCloudMusicApi 与 analytics 服务，再创建窗口；应用退出时关闭本地服务。渲染进程优先使用 preload 暴露的 apiPort，也支持通过 VITE_NETEASE_API_BASE_URL 配置远程或自定义 API 地址。")


def setup_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        try:
            style = styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8 if name == "Heading 3" else 12)
        style.paragraph_format.space_after = Pt(5)


def build_document():
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Missing converted template: {TEMPLATE_PATH}")
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    create_architecture_diagram(ASSET_DIR / "architecture.png")
    create_ui_flow_diagram(ASSET_DIR / "ui_flow.png")
    create_sequence_diagram(ASSET_DIR / "sequence_online_play.png")
    create_class_diagram(ASSET_DIR / "class_diagram.png")
    create_data_deployment_diagram(ASSET_DIR / "data_deployment.png")

    doc = Document(str(TEMPLATE_PATH))
    setup_styles(doc)
    update_cover(doc)
    rebuild_front_matter(doc)
    clear_body(doc)
    build_content(doc)
    doc.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
