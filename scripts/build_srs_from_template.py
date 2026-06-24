from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


TEMPLATE_PATH = Path("docs/template-srs-jasmine.docx")
OUTPUT_DOCX = Path("output/docx/Jasmine-软件需求规格说明书.docx")
ASSET_DIR = Path("output/docx/assets")
FONT_PATH = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
DATE_TEXT = "2026-06-22"


def get_font(size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_PATH, size=size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for ch in text:
        candidate = current + ch
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill):
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=8, align="center")
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    x = box[0] + (box[2] - box[0] - tw) / 2
    y = box[1] + (box[3] - box[1] - th) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=8, align="center")


def add_arrow(draw: ImageDraw.ImageDraw, x1: int, y1: int, x2: int, y2: int, fill: str, width: int = 3):
    draw.line((x1, y1, x2, y2), fill=fill, width=width)
    if x2 >= x1:
        pts = [(x2, y2), (x2 - 16, y2 - 8), (x2 - 16, y2 + 8)]
    else:
        pts = [(x2, y2), (x2 + 16, y2 - 8), (x2 + 16, y2 + 8)]
    draw.polygon(pts, fill=fill)


def create_use_case_diagram(output_path: Path):
    img = Image.new("RGB", (1600, 980), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(36)
    label_font = get_font(24)
    small_font = get_font(20)

    draw.text((70, 36), "Jasmine 系统用例图", fill="#123b63", font=title_font)

    boundary = (360, 120, 1290, 860)
    draw.rounded_rectangle(boundary, radius=18, outline="#7ea3c6", width=4)
    draw.text((boundary[0] + 24, boundary[1] + 16), "Jasmine 系统", fill="#3a5f83", font=label_font)

    def draw_actor(x: int, y: int, name: str):
        draw.ellipse((x - 30, y, x + 30, y + 60), outline="#34495e", width=4)
        draw.line((x, y + 60, x, y + 160), fill="#34495e", width=4)
        draw.line((x - 50, y + 95, x + 50, y + 95), fill="#34495e", width=4)
        draw.line((x, y + 160, x - 45, y + 220), fill="#34495e", width=4)
        draw.line((x, y + 160, x + 45, y + 220), fill="#34495e", width=4)
        name_box = (x - 90, y + 235, x + 90, y + 280)
        draw_centered_text(draw, name_box, name, label_font, "#2d3e50")

    draw_actor(140, 330, "用户")
    draw_actor(1460, 330, "网易云\n音乐服务")

    cases = [
        ("导入本地音乐", (540, 190, 860, 270)),
        ("搜索在线音乐", (900, 190, 1220, 270)),
        ("播放歌曲", (700, 330, 1020, 410)),
        ("管理歌单", (540, 470, 860, 550)),
        ("查看歌词", (900, 470, 1220, 550)),
        ("登录网易云", (700, 610, 1020, 690)),
        ("浏览推荐与排行榜", (700, 750, 1080, 830)),
    ]

    for name, box in cases:
        draw.ellipse(box, outline="#2d6aa0", width=4, fill="#f5fbff")
        draw_centered_text(draw, box, name, label_font, "#234f78")

    user_anchor = (190, 430)
    for _, box in cases:
        add_arrow(draw, user_anchor[0], user_anchor[1], box[0], (box[1] + box[3]) // 2, "#8ca5bc", width=2)

    service_anchor = (1410, 430)
    for index in [1, 4, 5, 6]:
        box = cases[index][1]
        add_arrow(draw, service_anchor[0], service_anchor[1], box[2], (box[1] + box[3]) // 2, "#8ca5bc", width=2)

    note = "说明：用户既可以管理本地音乐，也可以通过网易云服务完成搜索、登录、推荐访问与在线播放。"
    lines = wrap_text(draw, note, small_font, 1440)
    draw.multiline_text((80, 900), "\n".join(lines), font=small_font, fill="#4a6178", spacing=6)
    img.save(output_path)


def create_sequence_diagram_import(output_path: Path):
    img = Image.new("RGB", (1700, 1050), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34)
    label_font = get_font(22)
    small_font = get_font(18)

    draw.text((70, 30), "关键用例时序图 1：本地音乐导入", font=title_font, fill="#123b63")

    participants = [
        ("用户", 150),
        ("音乐库界面", 430),
        ("PlayerContext", 760),
        ("元数据解析服务", 1080),
        ("本地存储", 1420),
    ]
    top = 130
    bottom = 950
    for name, x in participants:
        draw.rounded_rectangle((x - 90, top, x + 90, top + 50), radius=12, outline="#6488aa", width=3, fill="#f7fbff")
        draw_centered_text(draw, (x - 86, top + 6, x + 86, top + 44), name, label_font, "#234f78")
        draw.line((x, top + 50, x, bottom), fill="#b8c9d8", width=2)

    steps = [
        (150, 430, 220, "选择音频文件"),
        (430, 760, 310, "调用 importFiles(files)"),
        (760, 1080, 400, "循环执行 extractMetadata(file)"),
        (1080, 760, 490, "返回歌曲名/歌手/时长等元数据"),
        (760, 1420, 580, "保存音频文件到 IndexedDB"),
        (760, 1420, 670, "保存歌曲元数据到 localStorage"),
        (760, 430, 760, "更新音乐库状态"),
        (430, 150, 850, "展示导入结果与歌曲列表"),
    ]
    for x1, x2, y, label in steps:
        add_arrow(draw, x1, y, x2, y, "#315d84", width=3)
        lines = wrap_text(draw, label, small_font, abs(x2 - x1) - 40)
        draw.multiline_text((min(x1, x2) + 20, y - 36), "\n".join(lines), font=small_font, fill="#3d556b", spacing=4)

    img.save(output_path)


def create_sequence_diagram_online(output_path: Path):
    img = Image.new("RGB", (1760, 1120), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34)
    label_font = get_font(22)
    small_font = get_font(18)

    draw.text((70, 30), "关键用例时序图 2：在线搜索并播放歌曲", font=title_font, fill="#123b63")

    participants = [
        ("用户", 150),
        ("发现页/搜索页", 450),
        ("neteaseApi", 780),
        ("PlayerContext", 1110),
        ("音频播放器", 1450),
    ]
    top = 130
    bottom = 1020
    for name, x in participants:
        draw.rounded_rectangle((x - 92, top, x + 92, top + 50), radius=12, outline="#6488aa", width=3, fill="#f7fbff")
        draw_centered_text(draw, (x - 88, top + 6, x + 88, top + 44), name, label_font, "#234f78")
        draw.line((x, top + 50, x, bottom), fill="#b8c9d8", width=2)

    steps = [
        (150, 450, 220, "输入关键词并发起搜索"),
        (450, 780, 310, "调用 searchOnline(keywords)"),
        (780, 450, 400, "返回歌曲/歌单/艺人结果"),
        (150, 450, 490, "点击目标歌曲"),
        (450, 780, 580, "请求音频地址与歌词"),
        (780, 450, 670, "返回 song/url 与 lyric"),
        (450, 1110, 760, "dispatch PLAY_SONG"),
        (1110, 1450, 850, "加载音频并调用 play()"),
        (1450, 1110, 940, "回传进度与播放状态"),
        (1110, 450, 1030, "同步当前歌曲/歌词/队列"),
    ]
    for x1, x2, y, label in steps:
        add_arrow(draw, x1, y, x2, y, "#315d84", width=3)
        lines = wrap_text(draw, label, small_font, max(140, abs(x2 - x1) - 40))
        draw.multiline_text((min(x1, x2) + 18, y - 34), "\n".join(lines), font=small_font, fill="#3d556b", spacing=4)

    img.save(output_path)


def create_prototype_wireframe(output_path: Path):
    img = Image.new("RGB", (1600, 940), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34)
    label_font = get_font(24)
    small_font = get_font(20)

    draw.text((70, 30), "Jasmine 原型结构示意", font=title_font, fill="#123b63")
    app = (120, 110, 1480, 850)
    draw.rounded_rectangle(app, radius=20, outline="#6d8ba9", width=4, fill="#fdfefe")

    sidebar = (150, 150, 410, 720)
    main = (450, 150, 1450, 720)
    player = (150, 740, 1450, 820)

    draw.rounded_rectangle(sidebar, radius=18, outline="#4f6880", width=3, fill="#15304a")
    draw.rounded_rectangle(main, radius=18, outline="#9db3c7", width=3, fill="#eef5fb")
    draw.rounded_rectangle(player, radius=14, outline="#4f6880", width=3, fill="#dce8f2")

    draw.text((185, 175), "侧边栏导航", font=label_font, fill="white")
    sidebar_items = ["发现", "我的音乐", "搜索", "设置", "下载", "用户歌单"]
    for i, item in enumerate(sidebar_items):
        y = 235 + i * 64
        draw.rounded_rectangle((182, y, 378, y + 42), radius=10, fill="#234a70", outline="#6b90b0")
        draw.text((205, y + 9), item, font=small_font, fill="#eaf4fc")

    draw.text((490, 175), "主内容区", font=label_font, fill="#234f78")
    cards = [
        (500, 230, 780, 360, "热门搜索 / 推荐歌单"),
        (810, 230, 1100, 360, "排行榜"),
        (1130, 230, 1410, 360, "艺人 / 专辑详情"),
        (500, 400, 1410, 520, "歌曲列表 / 歌单详情 / 历史记录"),
        (500, 560, 920, 680, "歌词 / 相似歌曲"),
        (950, 560, 1410, 680, "播放队列 / 收藏专辑"),
    ]
    for x1, y1, x2, y2, name in cards:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=16, outline="#90abc2", width=3, fill="white")
        draw_centered_text(draw, (x1 + 20, y1 + 20, x2 - 20, y2 - 20), name, small_font, "#45627d")

    draw.text((180, 762), "底部播放器：封面、歌曲信息、播放控制、音量、队列入口", font=small_font, fill="#24445f")
    img.save(output_path)


def set_run_font(run, name: str = "宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_paragraph(doc: Document, text: str, style: str = "Normal", bold_prefix: str | None = None, center: bool = False):
    p = doc.add_paragraph(style=style)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bold_prefix and text.startswith(bold_prefix):
        run1 = p.add_run(bold_prefix)
        run1.bold = True
        set_run_font(run1)
        rest = text[len(bold_prefix):]
        if rest:
            run2 = p.add_run(rest)
            set_run_font(run2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_blank(doc: Document):
    doc.add_paragraph("")


def add_image(doc: Document, path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.9))
    cap = doc.add_paragraph(style="Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def update_title_page(doc: Document):
    doc.paragraphs[0].text = "文档编号：Jasmine - SRS - V1.0"
    doc.paragraphs[11].text = "Jasmine"
    doc.paragraphs[12].text = "软件需求规格说明书"
    doc.paragraphs[31].text = f"日期：{DATE_TEXT}"


def update_change_table(doc: Document):
    table = doc.tables[0]
    row = table.rows[1].cells
    row[0].text = "1"
    row[1].text = DATE_TEXT
    row[2].text = "项目组"
    row[3].text = "依据 Jasmine 项目现状完成首版软件需求规格说明书整理"
    row[4].text = "V1.0"


def update_toc(doc: Document):
    entries = [
        "目录",
        "1. 引言",
        "1.1 编写目的",
        "1.2 读者对象",
        "1.3 软件项目概述",
        "1.4 文档概述",
        "1.5 定义",
        "1.6 参考资料",
        "2. 软件的一般性描述",
        "2.1 软件产品与其环境之间的关系",
        "2.2 限制与约束",
        "2.3 假设与前提条件",
        "3. 软件功能需求描述",
        "3.1 软件功能概述",
        "3.2 软件需求的用例模型",
        "3.3 软件需求的分析模型",
        "4. 其它软件需求描述",
        "4.1 性能要求",
        "4.2 设计约束",
        "4.3 界面要求",
        "4.4 进度要求",
        "4.5 交付要求",
        "4.6 验收要求",
        "5. 软件原型",
    ]
    doc.paragraphs[38].text = entries[0]
    for index, text in enumerate(entries[1:], start=39):
        doc.paragraphs[index].text = text


def clear_body(doc: Document):
    for paragraph in list(doc.paragraphs[63:])[::-1]:
        delete_paragraph(paragraph)


def build_content(doc: Document):
    add_paragraph(doc, "1. 引言", style="Heading 1")
    add_paragraph(doc, "1.1 编写目的", style="Heading 3")
    add_paragraph(doc, "本文档用于对 Jasmine 桌面音乐播放器项目的需求进行系统化说明，明确软件的功能边界、使用对象、关键用例、约束条件、交付要求与验收依据，为后续实现、测试、答辩展示和文档归档提供统一参考。", style="Normal")

    add_paragraph(doc, "1.2 读者对象", style="Heading 3")
    add_paragraph(doc, "本文档的主要读者包括项目开发者、课程指导教师、项目评审者、测试人员以及需要快速了解系统范围与核心能力的协作成员。对开发者而言，文档强调需求边界和模块关系；对评审者而言，文档强调系统完整性、关键场景与可验证性。", style="Normal")

    add_paragraph(doc, "1.3 软件项目概述", style="Heading 3")
    add_paragraph(doc, "项目名称：Jasmine。该项目是一个基于 Vite、React、TypeScript 与 Electron 构建的桌面音乐播放器，支持本地音乐导入、网易云在线搜索与播放、歌词展示、歌单管理、播放历史、国际化和桌面端安装发布。系统目标是在统一界面中完成本地资源管理与在线音乐发现两类核心任务。", style="Normal")

    add_paragraph(doc, "1.4 文档概述", style="Heading 3")
    add_paragraph(doc, "本文档首先介绍项目目标和外部环境，其后从功能需求、用例模型、分析模型和其它软件需求几个方面展开说明，最后补充原型结构描述。文档结构遵循软件需求规格说明书的常见写法，但内容结合 Jasmine 项目的实际实现状态进行了定制化整理。", style="Normal")

    add_paragraph(doc, "1.5 定义", style="Heading 3")
    add_paragraph(doc, "SRS：软件需求规格说明书，用于描述系统应提供的功能与约束。", style="Normal")
    add_paragraph(doc, "用例：从用户目标出发描述系统提供的完整交互场景。", style="Normal")
    add_paragraph(doc, "播放队列：系统当前播放上下文中的歌曲集合及其顺序信息。", style="Normal")
    add_paragraph(doc, "网易云服务：本项目所接入的在线音乐数据与播放地址来源。", style="Normal")
    add_paragraph(doc, "持久化存储：指 IndexedDB 与 localStorage 组成的本地数据保存方案。", style="Normal")

    add_paragraph(doc, "1.6 参考资料", style="Heading 3")
    add_paragraph(doc, "《Jasmine README》：项目功能、技术栈与运行方式说明。", style="Normal")
    add_paragraph(doc, "《CHANGELOG》：项目版本演进与已实现功能记录。", style="Normal")
    add_paragraph(doc, "《Jasmine 项目需求分析》：前期需求梳理文稿与结构化分析结论。", style="Normal")
    add_paragraph(doc, "NeteaseCloudMusicApi 文档：在线音乐接口集成参考。", style="Normal")

    add_paragraph(doc, "2. 软件的一般性描述", style="Heading 1")
    add_paragraph(doc, "2.1 软件产品与其环境之间的关系", style="Heading 3")
    add_paragraph(doc, "Jasmine 运行于桌面操作系统环境中，以 Electron 作为桌面封装层，以 React 页面作为交互界面，以 HTML5 Audio 作为播放引擎，并通过 NeteaseCloudMusicApi 服务访问在线音乐数据。系统还依赖本地文件系统提供用户导入的音频文件来源，并通过 IndexedDB 与 localStorage 保存业务数据。", style="Normal")
    add_paragraph(doc, "从外部关系上看，系统的主要交互对象包括：普通用户、桌面操作系统、本地文件资源、网易云音乐接口服务以及桌面安装运行环境。", style="Normal")

    add_paragraph(doc, "2.2 限制与约束", style="Heading 3")
    add_paragraph(doc, "系统前端采用 Vite + React + TypeScript 技术栈，桌面端采用 Electron 封装，因此界面组织、状态管理和桌面通信方案需与现有技术选型保持一致。", style="Normal")
    add_paragraph(doc, "在线搜索、歌词与在线播放能力依赖本地或远程部署的网易云 API 服务，因此在未启动该服务时，相关在线功能将受限。", style="Normal")
    add_paragraph(doc, "系统当前主要面向 macOS 与 Windows 环境发布，Linux 桌面版本支持尚不完整。", style="Normal")
    add_paragraph(doc, "由于项目处于毕业设计场景，需求覆盖范围应优先保证核心功能完整与演示效果清晰，而非追求企业级超大规模架构。", style="Normal")

    add_paragraph(doc, "2.3 假设与前提条件", style="Heading 3")
    add_paragraph(doc, "假设用户具备基本的桌面应用使用能力，并能够理解歌曲导入、歌单组织和在线播放等常见播放器操作。", style="Normal")
    add_paragraph(doc, "假设运行环境允许应用访问本地文件、网络接口和浏览器级音频播放能力。", style="Normal")
    add_paragraph(doc, "假设项目后续仍以当前代码仓库为唯一实现来源，需求分析中的模块命名和能力边界与现有实现保持一致。", style="Normal")

    add_paragraph(doc, "3. 软件功能需求描述", style="Heading 1")
    add_paragraph(doc, "3.1 软件功能概述", style="Heading 3")
    add_paragraph(doc, "Jasmine 的核心功能可以概括为六类：本地音乐导入与管理、在线音乐搜索与发现、歌曲播放控制、歌单与收藏管理、歌词与沉浸式播放体验、桌面端安装与持续使用支持。", style="Normal")
    add_paragraph(doc, "在用户视角下，系统既要支持“把本地歌放进来并长期管理”，也要支持“在线找到歌并立即播放”，同时保证播放状态、历史记录和界面偏好可以被长期保存。", style="Normal")

    add_paragraph(doc, "3.2 软件需求的用例模型", style="Heading 3")
    add_paragraph(doc, "3.2.1 参与者说明", style="Heading 3")
    add_paragraph(doc, "主要参与者为用户。用户既包括以本地音乐为主的个人使用者，也包括希望搜索在线内容并进行沉浸式播放的网络音乐用户。辅助参与者包括本地文件系统与网易云音乐服务，它们分别为本地导入和在线内容访问提供支持。", style="Normal")
    add_paragraph(doc, "3.2.2 用例图", style="Heading 3")
    add_image(doc, ASSET_DIR / "use_case_diagram.png", "图 3-1 Jasmine 系统用例图")

    add_paragraph(doc, "3.2.3 关键用例说明", style="Heading 3")
    add_paragraph(doc, "UC-01 本地音乐导入", style="Normal", bold_prefix="UC-01 本地音乐导入")
    add_paragraph(doc, "参与者：用户。前置条件：用户已准备本地音频文件。基本流程：用户在音乐库中选择导入文件，系统解析元数据，保存音频文件和歌曲信息，并刷新音乐库列表。后置结果：歌曲可在“全部歌曲”或“本地歌曲”中被查看、播放与收藏。", style="Normal")
    add_paragraph(doc, "UC-02 在线搜索并播放歌曲", style="Normal", bold_prefix="UC-02 在线搜索并播放歌曲")
    add_paragraph(doc, "参与者：用户、网易云音乐服务。前置条件：在线接口服务可访问。基本流程：用户输入关键词，系统返回歌曲、歌单、艺人和专辑结果；用户点击歌曲后系统获取音频地址与歌词并进入播放状态。后置结果：当前歌曲进入播放队列，歌词与进度同步更新。", style="Normal")
    add_paragraph(doc, "UC-03 歌单组织与收藏管理", style="Normal", bold_prefix="UC-03 歌单组织与收藏管理")
    add_paragraph(doc, "参与者：用户。前置条件：系统中已存在歌曲数据。基本流程：用户创建歌单、添加或移除歌曲、修改歌单名称、设置歌单封面，并使用“喜欢的音乐”或专辑收藏能力管理偏好内容。后置结果：用户个性化内容被持久化保存，并可在后续使用中直接恢复。", style="Normal")
    add_paragraph(doc, "UC-04 桌面端持续使用", style="Normal", bold_prefix="UC-04 桌面端持续使用")
    add_paragraph(doc, "参与者：用户、桌面操作系统。基本流程：用户安装并启动软件，系统恢复播放状态、语言偏好、历史记录和用户歌单。后置结果：播放器具备可持续使用的软件体验，而非一次性演示页面。", style="Normal")

    add_paragraph(doc, "3.3 软件需求的分析模型", style="Heading 3")
    add_paragraph(doc, "3.3.1 本地音乐导入时序图", style="Heading 3")
    add_image(doc, ASSET_DIR / "sequence_import.png", "图 3-2 本地音乐导入关键时序图")
    add_paragraph(doc, "该时序图体现了系统在导入本地音频时的核心调用关系：界面触发导入、播放器上下文统一协调、元数据服务解析文件信息、本地存储模块分别保存二进制音频与业务元数据。此过程强调系统的持久化能力和模块分工。", style="Normal")

    add_paragraph(doc, "3.3.2 在线搜索与播放时序图", style="Heading 3")
    add_image(doc, ASSET_DIR / "sequence_online.png", "图 3-3 在线搜索并播放关键时序图")
    add_paragraph(doc, "该时序图说明了在线搜索与播放的主链路：用户在搜索页触发检索，接口封装层请求网易云服务返回结果，播放上下文进一步获取音频地址并驱动播放器开始播放，同时同步歌词、进度与队列状态。", style="Normal")

    add_paragraph(doc, "3.3.3 核心分析对象说明", style="Heading 3")
    add_paragraph(doc, "从分析对象角度看，系统主要由界面层、播放器状态层、在线接口层、本地持久化层和桌面封装层构成。界面层负责与用户交互；播放器状态层负责播放队列、进度和控制命令；在线接口层负责搜索、推荐、歌词和音频地址获取；本地持久化层负责保存歌曲、歌单与历史记录；桌面封装层负责跨平台运行和资源访问。", style="Normal")

    add_paragraph(doc, "4. 其它软件需求描述", style="Heading 1")
    add_paragraph(doc, "4.1 性能要求", style="Heading 3")
    add_paragraph(doc, "系统在本地音乐加载、界面切换、歌词展示和常规在线播放操作中应保持良好的响应速度。对于大歌单加载、分页请求、音频地址获取等高延迟场景，系统应具备渐进加载与预取优化能力，避免用户长时间等待。", style="Normal")

    add_paragraph(doc, "4.2 设计约束", style="Heading 3")
    add_paragraph(doc, "系统需遵循既有技术架构：Vite 负责构建，React 负责界面，TypeScript 提供类型安全，Electron 提供桌面封装。在线能力依赖 NeteaseCloudMusicApi 服务，本地数据保存依赖 IndexedDB 与 localStorage。安全性方面应避免在渲染进程中直接暴露过高权限，可靠性方面应优先保证播放主链路和数据恢复能力。", style="Normal")

    add_paragraph(doc, "4.3 界面要求", style="Heading 3")
    add_paragraph(doc, "系统界面应采用统一的深色视觉风格，并确保导航结构清晰、播放器操作可直达、主要视图切换自然。核心界面至少应包括发现页、音乐库、搜索页、全屏播放页、歌单详情页和设置页。对桌面端而言，常用操作应尽量减少跳转层级，并提供稳定的底部播放器入口。", style="Normal")

    add_paragraph(doc, "4.4 进度要求", style="Heading 3")
    add_paragraph(doc, "系统当前已完成主要功能闭环，后续进度重点应放在三项任务：第一，修复大型在线歌单播放稳定性问题；第二，完善桌面播放器常用能力，例如设置中心与更强的队列管理；第三，逐步补充多音源接入、推荐解释和可视化分析等亮点功能。", style="Normal")

    add_paragraph(doc, "4.5 交付要求", style="Heading 3")
    add_paragraph(doc, "项目最终交付应至少包括：完整源代码仓库、需求与设计文档、可运行的桌面安装包或构建结果、必要的说明文档以及用于展示的需求规格说明书与分析文档。交付形式应支持电子文件归档，并便于课程评审和后续维护。", style="Normal")

    add_paragraph(doc, "4.6 验收要求", style="Heading 3")
    add_paragraph(doc, "系统验收应以核心功能可用为基础，主要包括：本地音乐导入成功、歌曲可正常播放、歌单与收藏可持久化保存、在线搜索和在线播放可完成、歌词显示可正常同步、桌面端可启动运行。对于严重缺陷，如大歌单无法播放，应作为重点核验项。", style="Normal")

    add_paragraph(doc, "5. 软件原型", style="Heading 1")
    add_paragraph(doc, "5.1 原型结构说明", style="Heading 3")
    add_paragraph(doc, "Jasmine 的当前原型不是静态页面集合，而是已经具备可运行能力的交互原型。其结构由侧边栏导航、主内容区域和底部播放器三部分组成：侧边栏负责视图切换和入口组织，主内容区负责发现、搜索和音乐库展示，底部播放器负责持续播放控制。", style="Normal")
    add_image(doc, ASSET_DIR / "prototype_wireframe.png", "图 5-1 Jasmine 原型结构示意")

    add_paragraph(doc, "5.2 主要界面说明", style="Heading 3")
    add_paragraph(doc, "发现页用于展示热门搜索、推荐歌单、排行榜和艺人/专辑详情；音乐库页用于呈现本地歌曲、自建歌单、收藏专辑和播放历史；搜索页用于本地检索；全屏播放页提供歌词、相似歌曲、分享和沉浸式体验；设置页当前主要负责语言切换，并为后续扩展系统设置预留入口。", style="Normal")


def build_document():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

    create_use_case_diagram(ASSET_DIR / "use_case_diagram.png")
    create_sequence_diagram_import(ASSET_DIR / "sequence_import.png")
    create_sequence_diagram_online(ASSET_DIR / "sequence_online.png")
    create_prototype_wireframe(ASSET_DIR / "prototype_wireframe.png")

    doc = Document(str(TEMPLATE_PATH))
    update_title_page(doc)
    update_change_table(doc)
    update_toc(doc)
    clear_body(doc)
    build_content(doc)
    doc.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
